#!/usr/bin/env python
# coding: utf-8

# In[10]:


#importing libraries
import pdfminer  # pdf2txt
import pandas as pd  # data manipulation
import re   # regular exression manipulation
import os   # operating system commands
import spacy  #NLP  "natural language processing"
import time  # just for the naming of the output CSV
from docx import Document


# In[2]:


import pdf2txt #import the file we downloaded from pdfminer repo


# In[8]:


# specifing the path for the files to be converted
path = 'D:\\ALL_data_science\\Freecodecamp\\Automation_with_python\\Resume_info_extractor\\output\\txt'
direct = os.listdir(path)


# In[14]:


# converting the files
for i in direct:
    document = Document()
    document.add_heading(i, 0)
    myfile = open('\\output\\txt'+i).read()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
    p = document.add_paragraph(myfile)
    document.save('\\output\\doc'+i+'.docx')


# In[ ]:




